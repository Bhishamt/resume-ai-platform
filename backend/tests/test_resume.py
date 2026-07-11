"""Test suite for Resume Management & Parsing Engine."""

import os
from pathlib import Path

import docx
import fitz
import pytest

from app.core.config import settings


@pytest.fixture
def temp_pdf(tmp_path) -> Path:
    """Create a valid dummy PDF file for testing."""
    file_path = tmp_path / "test_resume.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (50, 50),
        "Jane Smith\njane.smith@example.com\n(555) 019-2834\n"
        "Skills: Python, Javascript, React, Docker, SQL\n"
        "Education: Bachelor of Computer Science, University of California\n"
        "Experience: Software Engineer Intern, Tech Corp",
    )
    doc.save(str(file_path))
    doc.close()
    return file_path


@pytest.fixture
def temp_docx(tmp_path) -> Path:
    """Create a valid dummy DOCX file for testing."""
    file_path = tmp_path / "test_resume.docx"
    doc = docx.Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("jane.smith@example.com")
    doc.add_paragraph("(555) 019-2834")
    doc.add_paragraph("Skills: Python, Javascript, React, Docker, SQL")
    doc.add_paragraph(
        "Education: Bachelor of Computer Science, University of California"
    )
    doc.add_paragraph("Experience: Software Engineer Intern, Tech Corp")
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def temp_txt(tmp_path) -> Path:
    """Create an unsupported text file for testing."""
    file_path = tmp_path / "invalid_file.txt"
    with open(file_path, "w") as f:
        f.write("This is unsupported text format.")
    return file_path


@pytest.fixture
def other_user_auth(client):
    """Register and log in a second test user to test ownership authorization."""
    user_data = {
        "full_name": "Other User",
        "email": "other@example.com",
        "password": "SecureP@ss2",
    }
    client.post("/api/v1/auth/register", json=user_data)
    login_resp = client.post(
        "/api/v1/auth/login",
        json={"email": user_data["email"], "password": user_data["password"]},
    )
    token = login_resp.json()["data"]["access_token"]
    return {"Authorization": f"Bearer {token}"}


# ---------- File Upload Tests ----------


class TestResumeUpload:
    def test_upload_pdf_success(self, client, auth_headers, temp_pdf):
        with open(temp_pdf, "rb") as f:
            response = client.post(
                "/api/v1/resumes/upload",
                files={"file": ("test_resume.pdf", f, "application/pdf")},
                data={"title": "My PDF Resume"},
                headers=auth_headers,
            )
        assert response.status_code == 201
        data = response.json()
        assert data["success"] is True
        assert data["data"]["title"] == "My PDF Resume"
        assert data["data"]["upload_status"] == "success"
        assert "Jane Smith" in data["data"]["parsed_text"]
        assert "jane.smith@example.com" in data["data"]["parsed_text"]

        # Clean up created file
        storage_path = data["data"]["storage_path"]
        assert os.path.exists(storage_path)
        # Delete via service or manually to leave disk clean
        os.remove(storage_path)

    def test_upload_docx_success(self, client, auth_headers, temp_docx):
        with open(temp_docx, "rb") as f:
            response = client.post(
                "/api/v1/resumes/upload",
                files={
                    "file": (
                        "test_resume.docx",
                        f,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=auth_headers,
            )
        assert response.status_code == 201
        data = response.json()
        assert data["success"] is True
        assert data["data"]["title"] == "test_resume.docx"
        assert data["data"]["upload_status"] == "success"
        assert "Jane Smith" in data["data"]["parsed_text"]

        # Clean up
        storage_path = data["data"]["storage_path"]
        os.remove(storage_path)

    def test_upload_invalid_extension(self, client, auth_headers, temp_txt):
        with open(temp_txt, "rb") as f:
            response = client.post(
                "/api/v1/resumes/upload",
                files={"file": ("invalid_file.txt", f, "text/plain")},
                headers=auth_headers,
            )
        assert response.status_code == 400
        assert "Only PDF and DOCX" in response.json()["message"]

    def test_upload_invalid_mime(self, client, auth_headers, temp_pdf):
        """Rejects PDF with text/plain mime type."""
        with open(temp_pdf, "rb") as f:
            response = client.post(
                "/api/v1/resumes/upload",
                files={"file": ("test_resume.pdf", f, "text/plain")},
                headers=auth_headers,
            )
        assert response.status_code == 400
        assert "MIME type" in response.json()["message"]

    def test_upload_oversized_file(self, client, auth_headers, tmp_path, monkeypatch):
        # Temporarily mock MAX_UPLOAD_SIZE to 50 bytes for test
        monkeypatch.setattr(settings, "MAX_UPLOAD_SIZE", 50)
        file_path = tmp_path / "large.pdf"
        # Write 100 bytes to be larger than 50 bytes limit
        with open(file_path, "wb") as f:
            f.write(b"0" * 100)

        with open(file_path, "rb") as f:
            response = client.post(
                "/api/v1/resumes/upload",
                files={"file": ("large.pdf", f, "application/pdf")},
                headers=auth_headers,
            )
        assert response.status_code == 400
        assert "exceeds the limit" in response.json()["message"]


# ---------- Details, Delete, Replace & Ownership ----------


class TestResumeManagement:
    @pytest.fixture
    def uploaded_resume(self, client, auth_headers, temp_pdf):
        """Fixture to upload a resume and clean up afterwards."""
        with open(temp_pdf, "rb") as f:
            response = client.post(
                "/api/v1/resumes/upload",
                files={"file": ("test_resume.pdf", f, "application/pdf")},
                headers=auth_headers,
            )
        resume_data = response.json()["data"]
        yield resume_data
        # Clean up file if still exists
        path = resume_data["storage_path"]
        if os.path.exists(path):
            os.remove(path)

    def test_get_resumes_list(self, client, auth_headers, uploaded_resume):
        response = client.get("/api/v1/resumes", headers=auth_headers)
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert len(data["data"]["items"]) >= 1
        assert data["data"]["total"] >= 1

    def test_get_resume_details(self, client, auth_headers, uploaded_resume):
        response = client.get(
            f"/api/v1/resumes/{uploaded_resume['id']}",
            headers=auth_headers,
        )
        assert response.status_code == 200
        data = response.json()
        assert data["data"]["id"] == uploaded_resume["id"]
        assert "Jane Smith" in data["data"]["parsed_text"]

    def test_get_resume_details_unauthorized(
        self, client, other_user_auth, uploaded_resume
    ):
        """Verifies ownership check: other user cannot fetch User A's resume details."""
        response = client.get(
            f"/api/v1/resumes/{uploaded_resume['id']}",
            headers=other_user_auth,
        )
        assert response.status_code == 403
        assert "permission" in response.json()["message"]

    def test_update_resume_title(self, client, auth_headers, uploaded_resume):
        response = client.put(
            f"/api/v1/resumes/{uploaded_resume['id']}",
            json={"title": "Updated Title"},
            headers=auth_headers,
        )
        assert response.status_code == 200
        assert response.json()["data"]["title"] == "Updated Title"

    def test_replace_resume(self, client, auth_headers, uploaded_resume, temp_docx):
        old_path = uploaded_resume["storage_path"]
        with open(temp_docx, "rb") as f:
            response = client.post(
                f"/api/v1/resumes/{uploaded_resume['id']}/replace",
                files={
                    "file": (
                        "replaced.docx",
                        f,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=auth_headers,
            )
        assert response.status_code == 200
        data = response.json()
        assert data["data"]["id"] == uploaded_resume["id"]
        assert data["data"]["original_filename"] == "replaced.docx"

        # Verify old file was deleted from disk and new file exists
        assert not os.path.exists(old_path)
        new_path = data["data"]["storage_path"]
        assert os.path.exists(new_path)

        # Cleanup
        os.remove(new_path)

    def test_delete_resume(self, client, auth_headers, uploaded_resume):
        path = uploaded_resume["storage_path"]
        assert os.path.exists(path)

        # Delete request
        response = client.delete(
            f"/api/v1/resumes/{uploaded_resume['id']}",
            headers=auth_headers,
        )
        assert response.status_code == 200
        assert response.json()["success"] is True

        # Verify DB entry deleted
        details_resp = client.get(
            f"/api/v1/resumes/{uploaded_resume['id']}",
            headers=auth_headers,
        )
        assert details_resp.status_code == 404

        # Verify physical file deleted from storage
        assert not os.path.exists(path)


# ---------- Upload History Tests ----------


class TestUploadHistory:
    def test_get_history(self, client, auth_headers, temp_pdf):
        # 1. Perform upload
        with open(temp_pdf, "rb") as f:
            client.post(
                "/api/v1/resumes/upload",
                files={"file": ("test_resume.pdf", f, "application/pdf")},
                headers=auth_headers,
            )

        # 2. Get history
        response = client.get("/api/v1/resumes/history", headers=auth_headers)
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert len(data["data"]) >= 1
        assert data["data"][0]["action"] == "upload"

        # Clean up files created in test
        # (We can fetch the list and delete the files physically)
        list_resp = client.get("/api/v1/resumes", headers=auth_headers)
        for item in list_resp.json()["data"]["items"]:
            client.delete(f"/api/v1/resumes/{item['id']}", headers=auth_headers)
