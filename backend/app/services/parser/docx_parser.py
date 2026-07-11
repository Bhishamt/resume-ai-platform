import docx

from app.services.parser.base_parser import BaseParser, ParsedResumeData


class DOCXParser(BaseParser):
    """Resume parser implementation for DOCX documents."""

    def parse(self, file_path: str) -> ParsedResumeData:
        """Extract text from DOCX paragraphs and return parsed metadata."""
        raw_text = ""
        try:
            doc = docx.Document(file_path)
            paragraphs_text = [p.text for p in doc.paragraphs if p.text]

            # Also extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            tables_text.append(cell.text)

            raw_text = "\n".join(paragraphs_text + tables_text)
        except Exception:
            raw_text = ""

        # Extract contact details and sections using base helper
        extracted = self.extract_metadata(raw_text)

        return ParsedResumeData(
            name=extracted["name"],
            email=extracted["email"],
            phone=extracted["phone"],
            skills=extracted["skills"],
            education=extracted["education"],
            experience=extracted["experience"],
            projects=extracted["projects"],
            certifications=extracted["certifications"],
            raw_text=raw_text,
        )
